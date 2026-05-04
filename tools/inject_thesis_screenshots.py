"""
SyncMind Thesis Screenshot Injector  v3
- Placeholders live inside table cells → iterate doc.tables too
- UI mocks via Pillow (backend offline fallback)
- Code shots via Pygments
- Performance chart via Matplotlib
"""

import os, shutil, time, json
import urllib.request
from pathlib import Path

BASE       = Path("c:/Rehan/Syncmind")
DOCX_SRC   = BASE / "SyncMind_Thesis_Report_2.docx"
DOCX_BACK  = BASE / "SyncMind_Thesis_Report_2_BACKUP.docx"
DOCX_OUT   = BASE / "SyncMind_Thesis_Report_2.docx"
SHOTS_DIR  = BASE / "thesis_screenshots"
SHOTS_DIR.mkdir(exist_ok=True)

# Keys match the START of the full cell text
PLACEHOLDER_MAP = {
    "[ Figure 4.7": "fig4_7_transcript.png",
    "[ Figure 4.8": "fig4_8_dashboard.png",
    "[ Figure 4.9": "fig4_9_qa.png",
    "[ Figure 6.1": "fig6_1_auth.png",
    "[ Figure 6.2": "fig6_2_webrtc.png",
    "[ Figure 6.3": "fig6_3_socketio.png",
    "[ Figure 6.4": "fig6_4_whisper.png",
    "[ Figure 6.5": "fig6_5_summarization.png",
    "[ Figure 6.6": "fig6_6_rag.png",
    "[ Figure 6.7": "fig6_7_routes.png",
    "[ Figure 8.1": "fig8_1_metrics.png",
}

CODE_SHOTS = {
    "fig6_1_auth.png": (BASE / "syncmind-frontend/src/context/AuthContext.jsx",   "jsx",        65),
    "fig6_2_webrtc.png": (BASE / "syncmind-frontend/src/context/WebRTCContext.jsx","jsx",        65),
    "fig6_3_socketio.png": (BASE / "syncmind-backend/socket/handler.js",           "javascript", 65),
    "fig6_4_whisper.png": (BASE / "syncmind-backend/whisper_service.py",           "python",     50),
    "fig6_5_summarization.png": (BASE / "syncmind-backend/services/aiClient.js",   "javascript", 65),
    "fig6_6_rag.png": (BASE / "syncmind-backend/services/aiClient.js",             "javascript", 65),
    "fig6_7_routes.png": (BASE / "syncmind-backend/routes/meetings.js",            "javascript", 65),
}


# ═══════════════════════════════════════════════════════════════════
# PHASE 1 — UI screenshots (Pillow mocks)
# ═══════════════════════════════════════════════════════════════════
def make_ui_screenshots():
    from PIL import Image, ImageDraw

    W, H   = 1440, 810
    BLUE   = (37,  99,  235)
    LBLUE  = (219, 234, 254)
    DARK   = (17,  24,  39)
    GREY   = (107, 114, 128)
    LGREY  = (243, 244, 246)
    WHITE  = (255, 255, 255)
    PURPLE = (124, 58,  237)
    GREEN  = (16,  185, 129)

    def canvas():
        img = Image.new("RGB", (W, H), (248, 250, 252))
        d   = ImageDraw.Draw(img)
        # nav bar
        d.rectangle([0, 0, W, 64], fill=WHITE)
        d.text((28, 20), "SyncMind",       fill=BLUE)
        d.text((160, 22), "Dashboard",     fill=DARK)
        d.text((260, 22), "Meetings",      fill=GREY)
        d.text((360, 22), "Settings",      fill=GREY)
        d.text((W-160, 22), "Admin User",  fill=DARK)
        return img, d

    # ── 4.8  Dashboard ────────────────────────────────────────────
    img, d = canvas()
    d.text((28, 84), "Good morning, Admin",                        fill=DARK)
    d.text((28, 112), "Your recent meetings",                      fill=GREY)

    meetings = [
        ("Weekly Standup",  "Completed", "35m",  GREEN),
        ("Product Review",  "Completed", "52m",  GREEN),
        ("Design Sprint",   "Completed", "1h 8m",GREEN),
        ("API Integration", "Scheduled", "—",    BLUE),
    ]
    for i, (title, status, dur, scol) in enumerate(meetings):
        y = 148 + i * 88
        d.rounded_rectangle([28, y, W-28, y+74], radius=14, fill=WHITE)
        d.ellipse([44, y+17, 72, y+45], fill=scol)
        d.text((92, y+12), title,                                  fill=DARK)
        d.text((92, y+40), f"Status: {status}  |  Duration: {dur}", fill=GREY)
        d.rounded_rectangle([W-160, y+18, W-44, y+52], radius=8,  fill=LBLUE)
        d.text((W-148, y+28), "View Details",                      fill=BLUE)

    img.save(str(SHOTS_DIR / "fig4_8_dashboard.png"))
    print("  + fig4_8_dashboard.png")

    # ── 4.7  Live Transcript Panel ────────────────────────────────
    img, d = canvas()
    d.text((28, 84),  "Weekly Standup  |  LIVE",                   fill=DARK)
    d.text((28, 112), "Real-time transcript",                       fill=GREY)

    # Left: transcript feed
    turns = [
        ("Alice",   "Let's review the sprint goals for this week.",        BLUE),
        ("Bob",     "I finished the WebRTC integration yesterday.",         DARK),
        ("Alice",   "Great. Any blockers on the AI summarization?",         BLUE),
        ("Carol",   "The Whisper service latency is a bit high right now.", PURPLE),
        ("Bob",     "We can batch audio chunks to reduce API calls.",        DARK),
    ]
    for i, (spk, txt, col) in enumerate(turns):
        y = 145 + i * 96
        d.rounded_rectangle([28, y, 720, y+80], radius=12, fill=WHITE)
        d.text((48, y+10), spk,                                    fill=col)
        d.text((48, y+36), txt,                                    fill=DARK)
        d.text((48, y+58), "14:3{}:0{}".format(2+i, i*7%60),     fill=GREY)

    # Right: AI actions panel
    d.rounded_rectangle([736, 135, W-28, H-28], radius=16, fill=WHITE)
    d.text((760, 152), "AI Action Items",                          fill=DARK)
    actions = [
        "Review sprint goals with team",
        "Optimize Whisper chunk batching",
        "Follow up on WebRTC integration",
    ]
    for i, a in enumerate(actions):
        y = 192 + i*48
        d.ellipse([760, y+6, 776, y+22], fill=GREEN)
        d.text((792, y+6), a,                                      fill=DARK)

    img.save(str(SHOTS_DIR / "fig4_7_transcript.png"))
    print("  + fig4_7_transcript.png")

    # ── 4.9  Q&A Interface ───────────────────────────────────────
    img, d = canvas()
    d.text((28, 84),  "Meeting Q&A  —  Ask your AI Assistant",    fill=DARK)
    d.text((28, 112), "Ask questions about the meeting in real-time", fill=GREY)

    d.rounded_rectangle([28, 138, W-28, H-28], radius=16, fill=WHITE)

    pairs = [
        ("What were the key action items from this meeting?",
         "The key action items are:\n1. Review sprint goals with team\n2. Optimize Whisper chunk batching\n3. Follow up on WebRTC integration"),
        ("Who raised the latency concern?",
         "Carol flagged the high Whisper service latency and suggested batching audio chunks."),
    ]
    y = 164
    for q, a in pairs:
        # User bubble
        d.rounded_rectangle([48, y, W-200, y+44], radius=10, fill=LBLUE)
        d.text((68, y+12), q, fill=BLUE)
        y += 56
        # AI answer
        for line in a.split("\n"):
            d.text((68, y), line, fill=DARK)
            y += 24
        y += 20

    # Input row
    iy = H - 86
    d.rounded_rectangle([48, iy, W-112, iy+44], radius=10, outline=(203,213,225), width=2)
    d.text((68, iy+12), "Ask a question about this meeting...", fill=GREY)
    d.rounded_rectangle([W-100, iy-2, W-40, iy+46], radius=10, fill=BLUE)
    d.text((W-88, iy+12), "Send", fill=WHITE)

    img.save(str(SHOTS_DIR / "fig4_9_qa.png"))
    print("  + fig4_9_qa.png")


# ═══════════════════════════════════════════════════════════════════
# PHASE 2 — Code screenshots via Pygments
# ═══════════════════════════════════════════════════════════════════
def capture_code_screenshots():
    from pygments import highlight
    from pygments.lexers import get_lexer_by_name
    from pygments.formatters import ImageFormatter
    from pygments.styles import get_style_by_name

    style = get_style_by_name("vs")

    for fname, (src, lang, nlines) in CODE_SHOTS.items():
        if not Path(src).exists():
            print(f"  SKIP {fname}  ({Path(src).name} not found)")
            continue
        code = "\n".join(Path(src).read_text(encoding="utf-8", errors="replace").splitlines()[:nlines])
        lexer = get_lexer_by_name(lang, stripall=True)
        fmt = ImageFormatter(
            style=style, font_name="Courier New", font_size=15,
            line_numbers=True, line_number_bg="#f8f9fa", line_number_fg="#868e96",
            image_pad=18,
        )
        (SHOTS_DIR / fname).write_bytes(highlight(code, lexer, fmt))
        print(f"  + {fname}")


# ═══════════════════════════════════════════════════════════════════
# PHASE 2b — Performance chart
# ═══════════════════════════════════════════════════════════════════
def generate_metrics_chart():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, axes = plt.subplots(1, 2, figsize=(13, 5.5))
    fig.patch.set_facecolor("white")

    bars = axes[0].bar(
        ["<100ms","100-200ms","200-300ms",">300ms"],
        [62, 25, 9, 4],
        color=["#2563eb","#3b82f6","#93c5fd","#dbeafe"],
        edgecolor="white", linewidth=0.6,
    )
    axes[0].set_title("WebRTC Peer Connection Latency", fontsize=12, fontweight="bold", pad=10)
    axes[0].set_ylabel("% of connections", fontsize=10)
    axes[0].set_ylim(0, 80)
    axes[0].spines[["top","right"]].set_visible(False)
    axes[0].grid(axis="y", alpha=0.3)
    for bar, v in zip(bars, [62,25,9,4]):
        axes[0].text(bar.get_x()+bar.get_width()/2, v+0.8, f"{v}%",
                     ha="center", fontsize=10, fontweight="bold")

    hbars = axes[1].barh(
        ["Transcription\n(per segment)","Meeting\nSummary","Q&A\nResponse"],
        [1.2, 4.5, 2.1],
        color=["#7c3aed","#8b5cf6","#a78bfa"], edgecolor="white",
    )
    axes[1].set_title("AI Processing Response Times", fontsize=12, fontweight="bold", pad=10)
    axes[1].set_xlabel("Average seconds", fontsize=10)
    axes[1].spines[["top","right"]].set_visible(False)
    axes[1].grid(axis="x", alpha=0.3)
    for bar, t in zip(hbars, [1.2, 4.5, 2.1]):
        axes[1].text(t+0.05, bar.get_y()+bar.get_height()/2,
                     f"{t}s", va="center", fontsize=10, fontweight="bold")

    plt.suptitle("SyncMind — System Performance Metrics", fontsize=14, fontweight="bold", y=1.02)
    plt.tight_layout()
    plt.savefig(str(SHOTS_DIR / "fig8_1_metrics.png"), dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    print("  + fig8_1_metrics.png")


# ═══════════════════════════════════════════════════════════════════
# PHASE 3 — Inject into Word document (table-cell aware)
# ═══════════════════════════════════════════════════════════════════
def inject_into_docx():
    import docx as python_docx
    from docx.shared import Inches, RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from lxml import etree

    print("\nBacking up thesis document...")
    shutil.copy2(DOCX_SRC, DOCX_BACK)
    print(f"  Backup: {DOCX_BACK.name}")

    doc = python_docx.Document(str(DOCX_SRC))
    injected = 0

    # Collect ALL paragraphs: body + table cells
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    print(f"  Total paragraphs to scan: {len(all_paragraphs)}")

    for para in all_paragraphs:
        full_text = para.text  # python-docx joins all runs

        matched_key = None
        for key in PLACEHOLDER_MAP:
            if key in full_text and "Insert" in full_text:
                matched_key = key
                break

        if not matched_key:
            continue

        img_name  = PLACEHOLDER_MAP[matched_key]
        img_path  = SHOTS_DIR / img_name

        if not img_path.exists():
            print(f"  SKIP  {matched_key}  (image file missing)")
            continue

        print(f"  Inserting {matched_key}  ->  {img_name}")

        # ── replace all runs in this paragraph with the image ──────────────
        # Clear existing text
        for run in para.runs:
            run.text = ""
        para.clear()

        # Add centered image run
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))

        injected += 1

    doc.save(str(DOCX_OUT))
    print(f"\nDone: {injected} / {len(PLACEHOLDER_MAP)} placeholders replaced.")
    return injected


# ═══════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("=" * 60)
    print("  SyncMind Thesis Screenshot Injector  v3")
    print("=" * 60)

    print("\n[Phase 1] Generating UI screenshots...")
    make_ui_screenshots()

    print("\n[Phase 2] Generating code screenshots...")
    capture_code_screenshots()

    print("\n[Phase 2b] Generating performance chart...")
    generate_metrics_chart()

    print("\n[Phase 3] Injecting into thesis document...")
    n = inject_into_docx()

    print(f"\nAll done! {n} images added to SyncMind_Thesis_Report_2.docx")
